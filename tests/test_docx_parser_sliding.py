"""Phase 8.0 —— 兜底滑窗切分单元测试。

覆盖：
1. 短文档（< 阈值）→ 单 _intro chunk，向后兼容（命名沿用，不改成 _full）
2. 长文档（≥ 阈值）→ 多 _window_N chunk
3. overlap 段落只复制文本、不复制图片（同图不出现在相邻 chunk）
4. 表格作为整段加入 buffer，不被切断
5. 单段超过 size 时整段保留，不切碎
6. 集成：构造一份长 FAQ docx，验证 parse_docx 切出 _window_N

> PLAN §四.3 原方案使用 `_full` 命名长文档兜底前的短文档分支；为满足 §九
> "现有 agv_demo / ifs_docs chunks.jsonl 不变" 的兼容性验收，本期保留 `_intro`
> 命名，仅长文档新增 `_window_N` 命名。
"""

from __future__ import annotations

from pathlib import Path
from typing import List

import pytest

from custom_app.services.docx_parser import (
    SLIDING_WINDOW_OVERLAP_CHARS,
    SLIDING_WINDOW_SIZE_CHARS,
    SLIDING_WINDOW_THRESHOLD_CHARS,
    _sliding_window_chunks,
    parse_docx,
)


# ─────────────────────────────────────────────────────────────────────────────
# 纯函数：_sliding_window_chunks
# ─────────────────────────────────────────────────────────────────────────────


class TestSlidingWindowPureFunction:
    """对 _sliding_window_chunks 的行为做最小化验证。"""

    def test_single_short_part_returns_one_window(self) -> None:
        """单段短文本 → 1 个 window。"""
        parts = ["这是一段短文本"]
        imgs: List[List[str]] = [[]]
        out = _sliding_window_chunks(parts, imgs, size=800, overlap=100)
        assert len(out) == 1
        assert out[0][0] == ["这是一段短文本"]
        assert out[0][1] == []

    def test_long_sequence_splits_into_multiple_windows(self) -> None:
        """多段总长远超 size → 拆成多个 window。"""
        # 6 段各 300 字 = 1800 字，size=800 → 至少 3 个 window
        parts = ["甲" * 300, "乙" * 300, "丙" * 300, "丁" * 300, "戊" * 300, "己" * 300]
        imgs: List[List[str]] = [[] for _ in parts]
        out = _sliding_window_chunks(parts, imgs, size=800, overlap=100)
        assert len(out) >= 3, f"expected ≥3 windows, got {len(out)}"
        # 每个 window 至少包含一段
        for lines, _ in out:
            assert lines

    def test_overlap_copies_text_but_not_images(self) -> None:
        """overlap 段落带图：图片只归原 chunk，不复制到下一 chunk。"""
        # 段 A=500 字 + 段 B=500 字（带 1 张图）+ 段 C=500 字（带 1 张图）
        # 累积顺序：A(500) → 加 B 后 1000>800 触发 flush，buf=[A,B], imgs=[imgB]
        # overlap=100，B(500) > 100 → tail 仅留 B；imgs 清空
        # 再加 C：buf=[B,C], imgs=[imgC]，B 的图 imgB 不在第二个 window
        parts = ["甲" * 500, "乙" * 500, "丙" * 500]
        imgs = [[], ["images/doc/imgB.png"], ["images/doc/imgC.png"]]
        out = _sliding_window_chunks(parts, imgs, size=800, overlap=100)
        assert len(out) >= 2

        all_imgs_per_window = [w[1] for w in out]
        flat = [p for w in all_imgs_per_window for p in w]
        # 每张图最多出现一次（imgB 和 imgC 各 1 次）
        assert flat.count("images/doc/imgB.png") == 1
        assert flat.count("images/doc/imgC.png") == 1

    def test_oversized_single_part_is_split_at_punctuation(self) -> None:
        """Phase 11.2 行为变更：单段超 size 时按句号切分，不再整段保留。

        旧版（Phase 8.0）"整段保留"策略导致 ifs_docs 887 字 chunk 把
        '出库类型 / 入库类型 / 品牌 / 零件状态 / 计划人' 5 个子项埋一起，
        '计划人'子项被 LLM 漏掉。新版必须按 separators 降级切到 ≤ size。
        """
        # 5 句各 300 字 + 句号 = 单段 1500+ 字
        sentences = "。".join(["庞" * 300] * 5) + "。"
        parts = [sentences]
        imgs: List[List[str]] = [[]]
        out = _sliding_window_chunks(parts, imgs, size=800, overlap=100)
        # 必须切出 ≥2 个 window（按 "。" 切）
        assert len(out) >= 2
        # 每个 window 都不应再超 size 太多（句末切，允许超少量给最后一句）
        for lines, _ in out:
            joined = "\n".join(lines)
            # 允许最长不超 size * 1.5（取决于句子长度）
            assert len(joined) <= 800 * 1.5

    def test_misaligned_inputs_raise(self) -> None:
        """parts 与 imgs_per_part 长度不一致 → ValueError。"""
        with pytest.raises(ValueError):
            _sliding_window_chunks(["a", "b"], [[]], size=800, overlap=100)

    def test_markdown_table_protected_from_split(self) -> None:
        """Markdown 表格（含分隔行）应作为整体保留，不被切到 size 边界。

        _PROTECTED_PATTERNS 含表格行/分隔行模式；新版 _split_text_recursive
        会把 protected 区段整体保留（即便超 size，宁可单 chunk 大也不破表格）。
        """
        # 构造标准 markdown 表格（带 |---| 分隔行）
        table_md = (
            "| 列1 | 列2 | 列3 |\n"
            "| --- | --- | --- |\n"
            + "\n".join([f"| 数据{i} | 数据{i} | 数据{i} |" for i in range(20)])
        )
        # 普通前后文本 + 中间表格
        parts = ["前置说明文字段落甲" * 80, table_md, "尾部说明文字段落乙" * 80]
        imgs: List[List[str]] = [[] for _ in parts]
        out = _sliding_window_chunks(parts, imgs, size=400, overlap=80)
        # 表格分隔行 "| --- | --- | --- |" 不应被切断
        all_text = "\n".join(ln for w_lines, _ in out for ln in w_lines)
        assert "| --- | --- | --- |" in all_text, "表格分隔行应被保护不切断"
        # 表格的数据行也应完整存在
        assert "| 数据10 | 数据10 | 数据10 |" in all_text

    def test_img_placeholder_protected_from_split(self) -> None:
        """[IMG: path] 占位作为 protected 模式，不应被切到中间。"""
        # 前文 800 字 + [IMG: ...] + 后文
        big_text = "甲" * 800
        img_placeholder = "[IMG: images/foo/img_001.png]"
        parts = [big_text, img_placeholder, "尾部"]
        imgs: List[List[str]] = [[], ["images/foo/img_001.png"], []]
        out = _sliding_window_chunks(parts, imgs, size=400, overlap=80)
        all_text = "\n".join(ln for w_lines, _ in out for ln in w_lines)
        # [IMG: ...] 必须完整出现，路径不被截断
        assert img_placeholder in all_text


# ─────────────────────────────────────────────────────────────────────────────
# 集成：parse_docx 走兜底路径
# ─────────────────────────────────────────────────────────────────────────────


def _make_long_faq_docx(tmp_path: Path, *, total_chars: int = 1500) -> Path:
    """构造一份既无 STEP 也无 Heading 的长文档（走 parse_docx 兜底）。"""
    from docx import Document  # type: ignore

    doc = Document()
    # 用 N 段 ~ 200 字普通段落（默认 Normal 样式，不会被 _paragraph_heading_label 识别）
    per_para = 200
    n = max(1, total_chars // per_para)
    for i in range(n):
        doc.add_paragraph(f"问答条目 {i + 1}：" + "内容字符" * (per_para // 4))
    path = tmp_path / "faq_long.docx"
    doc.save(str(path))
    return path


def _make_short_faq_docx(tmp_path: Path) -> Path:
    from docx import Document  # type: ignore

    doc = Document()
    doc.add_paragraph("这是一段非常短的 FAQ。")
    doc.add_paragraph("第二段也很短。")
    path = tmp_path / "faq_short.docx"
    doc.save(str(path))
    return path


class TestParseDocxFallbackRouting:
    """parse_docx 兜底分支按字符阈值路由。"""

    def test_short_document_keeps_single_intro_chunk(self, tmp_path: Path) -> None:
        """< 阈值的短文档保持 _intro 单 chunk（向后兼容现有 KB）。"""
        kb_root = tmp_path / "kb"
        kb_root.mkdir()
        docx = _make_short_faq_docx(tmp_path)
        chunks = parse_docx(docx, kb_root)
        assert len(chunks) == 1
        assert chunks[0]["id"].endswith("_intro")
        # 不应触发滑窗
        assert "_window_" not in chunks[0]["id"]

    def test_long_document_routed_to_sliding_windows(self, tmp_path: Path) -> None:
        """≥ 阈值的长文档切出多个 _window_N chunk。"""
        kb_root = tmp_path / "kb"
        kb_root.mkdir()
        # 1500 字明显 > 阈值 800，应至少切 2 块
        docx = _make_long_faq_docx(tmp_path, total_chars=1500)
        chunks = parse_docx(docx, kb_root)

        window_ids = [c["id"] for c in chunks if "_window_" in c["id"]]
        intro_ids = [c["id"] for c in chunks if c["id"].endswith("_intro")]

        assert len(window_ids) >= 2, f"expected ≥2 _window_N chunks, got: {[c['id'] for c in chunks]}"
        assert not intro_ids, "long document should not collapse to a single _intro chunk"

        # schema 字段保留（新 schema 兼容）
        for c in chunks:
            assert c["source_type"] == "sop_docx"
            assert c["parser"] == "docx_parser"
            assert isinstance(c["images"], list)

    def test_threshold_constants_are_sane(self) -> None:
        """阈值常量自洽：threshold/size 至少 ≥ overlap*2。"""
        assert SLIDING_WINDOW_THRESHOLD_CHARS >= SLIDING_WINDOW_OVERLAP_CHARS * 2
        assert SLIDING_WINDOW_SIZE_CHARS >= SLIDING_WINDOW_OVERLAP_CHARS * 2


# ─────────────────────────────────────────────────────────────────────────────
# Phase 11.2: 多并列子项压一段时能被正确切开（"计划人" bug 回归测）
# ─────────────────────────────────────────────────────────────────────────────


class TestPhase11_2FineGrainedSplit:
    """ifs_docs 'section_2' 887 字含多个并列子项的切分回归。"""

    def test_multi_subitems_in_one_paragraph_split_into_separate_chunks(self) -> None:
        """模拟 ifs_docs '库存基础数据_section_2' 把 5 个子项压一段的情况，
        应被切成多个 chunk，'计划人' 子项要能独立出现在某一 chunk 里。
        """
        # 模拟原始 ifs_docs 887 字内容结构
        section_2 = (
            "库存基础数据。"
            "出库类型：对产品手工出库原因的归类说明，库存件下发时调用此数据。"
            "操作：选择 Issue Type 页签，双击行或点击新建按钮，录入编码描述备注，保存。"
            "其中：出库类型编码不可修改，描述及备注可修改。"
            "入库类型：对产品手工入库原因的归类说明，库存件接收时调用此数据。"
            "操作：选择 Receive Type 页签，双击行或点击新建按钮，录入编码描述备注，保存。"
            "品牌 Product Code：用于划分产品属性，在库存件设置中使用该参数。"
            "操作：选择 Product Code 页签，双击行或点击新建按钮录入。"
            "零件状态 Part Status：定义库存件数量、需求、供应允许情况。"
            "操作：选择 Part Status 页签，可设置不允许现有量、不允许需求、不允许采购。"
            "计划人 Planners：指编码提供者或审核者，中国公司固定为编码审核人。"
            "前置条件：设置计划人前，必须先在人员中为计划人创建人员记录，否则无法调用。"
            "操作：选择 Planners 页签，双击行或新建按钮，通过 F8 选择人员，保存。"
        )
        # 用单段 parts 输入（模拟 ifs_docs section 是一段长文本）
        parts = [section_2]
        imgs: List[List[str]] = [[]]
        # 用 size=300 模拟更细粒度（默认 400 时 448 字段会切 2 块）
        out = _sliding_window_chunks(parts, imgs, size=300, overlap=80)
        # 必须切成 ≥2 个 chunk（旧版会塞 1 个 chunk）
        assert len(out) >= 2, f"expected ≥2 chunks, got {len(out)}"

        # 至少有一个 chunk 应明确讲"计划人"配置（Planners 页签 + F8 操作）
        chunks_text = ["\n".join(lines) for lines, _ in out]
        planner_chunks = [c for c in chunks_text if "计划人" in c and "Planners 页签" in c]
        assert planner_chunks, f"'计划人' 应独立出现在某 chunk: {chunks_text}"

        # 每个 chunk 都不应超 size * 1.5（避免退化到旧版整段保留）
        for c in chunks_text:
            assert len(c) <= 300 * 1.5, f"chunk too big: {len(c)} chars"
