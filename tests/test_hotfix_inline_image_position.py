"""
Hotfix TDD：保留 docx 内图片的原始位置

修复：当前 pack_chunk 把整段文本拼好后，把所有图片一股脑塞到文本末尾的
[IMAGES] 区。LLM 看到的就是"5 段说明 + 5 张图集中堆末尾"，输出时也只能这么排。

期望：chunks.jsonl 里的 contents 字段应在每个图片本来出现的位置插入 [IMG: <相对路径>]
内联标记，让 LLM 看到"步骤说明文字 [IMG: img_001] 步骤说明文字 [IMG: img_002]"
的真实顺序。

list_knowledge_chunks 工具返回时，把 [IMG: <相对路径>] 转换成已编码的 markdown
图片占位符（或保留供 LLM 替换），并且 image_urls 字段保持完整数组（兼容老代码）。
"""
from __future__ import annotations

import json
import re
from pathlib import Path

import pytest


# ─────────────────────────────────────────────────────────────────────────────
# parse_docx 输出的 contents 字段：图片应内联，而不是集中在 [IMAGES] 区
# ─────────────────────────────────────────────────────────────────────────────

class TestDocxParserInlineImages:
    """图片位置应保留在 contents 字段对应位置。"""

    def _make_minimal_docx(self, tmp_path: Path) -> Path:
        """构造一个含 'A → image1 → B → image2' 顺序的 docx。"""
        from docx import Document
        from docx.shared import Inches

        doc = Document()
        doc.add_paragraph("段落 A：第一段说明文字。")

        # 真正最小的 1x1 RGB PNG（69 字节），python-docx 能解析
        png_bytes = bytes.fromhex(
            "89504e470d0a1a0a0000000d4948445200000001000000010802000000907753de"
            "0000000c49444154789c63f8ffff3f0005fe02fe0def46b80000000049454e44ae426082"
        )
        img1_path = tmp_path / "img1.png"
        img1_path.write_bytes(png_bytes)
        img2_path = tmp_path / "img2.png"
        img2_path.write_bytes(png_bytes)

        # 段落 A → image1 → 段落 B → image2
        para = doc.add_paragraph()
        para.add_run().add_picture(str(img1_path), width=Inches(1))

        doc.add_paragraph("段落 B：第二段说明文字。")

        para = doc.add_paragraph()
        para.add_run().add_picture(str(img2_path), width=Inches(1))

        out = tmp_path / "test.docx"
        doc.save(str(out))
        return out

    def test_contents_has_inline_image_markers(self, tmp_path):
        from custom_app.services.docx_parser import parse_docx
        docx = self._make_minimal_docx(tmp_path)
        kb_root = tmp_path / "kb"

        chunks = parse_docx(docx, kb_root)
        assert chunks, "should produce at least one chunk"
        chunk = chunks[0]
        contents = chunk["contents"]

        # 不再有"集中堆末尾"的 [IMAGES] 区段
        assert "[IMAGES]" not in contents, (
            f"contents should not collect images at the end, got:\n{contents}"
        )
        # contents 中应有内联 [IMG: ...] 标记
        markers = re.findall(r"\[IMG:\s*([^\]]+)\]", contents)
        assert len(markers) >= 2, f"expected ≥2 inline markers, got: {markers}"

    def test_inline_markers_in_correct_order(self, tmp_path):
        """段落 A 的图必须排在段落 B 之前。"""
        from custom_app.services.docx_parser import parse_docx
        docx = self._make_minimal_docx(tmp_path)
        kb_root = tmp_path / "kb"

        chunks = parse_docx(docx, kb_root)
        contents = chunks[0]["contents"]

        idx_a = contents.find("段落 A")
        idx_img1 = contents.find("[IMG:")
        idx_b = contents.find("段落 B")
        idx_img2 = contents.rfind("[IMG:")

        assert idx_a < idx_img1 < idx_b < idx_img2, (
            f"expected A < img1 < B < img2, got positions: "
            f"A={idx_a} img1={idx_img1} B={idx_b} img2={idx_img2}\n"
            f"contents:\n{contents}"
        )

    def test_images_array_still_populated(self, tmp_path):
        """images 字段保持向后兼容（同时含 image_urls 用的源数据）。"""
        from custom_app.services.docx_parser import parse_docx
        docx = self._make_minimal_docx(tmp_path)
        kb_root = tmp_path / "kb"

        chunks = parse_docx(docx, kb_root)
        chunk = chunks[0]
        assert isinstance(chunk["images"], list)
        assert len(chunk["images"]) >= 2


# ─────────────────────────────────────────────────────────────────────────────
# list_knowledge_chunks 工具：把 [IMG: ...] 转换成已编码 URL 的 markdown 占位
# ─────────────────────────────────────────────────────────────────────────────

class TestListChunksInlineMarkdown:
    """list_chunks 返回的 contents 应已经把 [IMG: 路径] 渲染为可用的 markdown 占位。"""

    def test_inline_marker_becomes_markdown_image(self):
        from custom_app.services.tools.list_chunks import ListChunksTool
        rows = [{
            "id": "0",
            "title": "STEP 1",
            "doc": "IFS 系统培训手册--库存销售基础数据",
            "contents": (
                "出库类型说明文字。\n"
                "[IMG: images/IFS 系统培训手册--库存销售基础数据/img_0001.png]\n"
                "入库类型说明文字。\n"
                "[IMG: images/IFS 系统培训手册--库存销售基础数据/img_0002.png]"
            ),
            "images": [
                "images/IFS 系统培训手册--库存销售基础数据/img_0001.png",
                "images/IFS 系统培训手册--库存销售基础数据/img_0002.png",
            ],
        }]
        tool = ListChunksTool(rows=rows)
        result = tool.run(doc_id="IFS 系统培训手册--库存销售基础数据")
        assert len(result) == 1
        rendered = result[0]["contents"]

        # [IMG: ...] 标记应被替换为 markdown 图片，URL 已编码
        assert "[IMG:" not in rendered, (
            f"raw [IMG: ...] should be replaced with markdown image, got:\n{rendered}"
        )
        assert "![](" in rendered
        # 空格变 %20
        assert "%20" in rendered
        # 图片穿插（出库说明在 img_0001 前，入库在 img_0002 前）
        idx_out = rendered.find("出库类型")
        idx_img1 = rendered.find("img_0001")
        idx_in = rendered.find("入库类型")
        idx_img2 = rendered.find("img_0002")
        assert idx_out < idx_img1 < idx_in < idx_img2, (
            f"expected interleaved order, got: "
            f"出库={idx_out} img1={idx_img1} 入库={idx_in} img2={idx_img2}\n"
            f"contents:\n{rendered}"
        )

    def test_image_urls_field_unchanged(self):
        """新增的内联渲染不影响 image_urls 字段（向后兼容）。"""
        from custom_app.services.tools.list_chunks import ListChunksTool
        rows = [{
            "id": "0", "title": "T", "doc": "D",
            "contents": "x",
            "images": ["images/D/img_0001.png"],
        }]
        tool = ListChunksTool(rows=rows)
        result = tool.run(doc_id="D")
        assert "image_urls" in result[0]
        assert result[0]["image_urls"] == ["/images/D/img_0001.png"]

    def test_no_inline_markers_no_change(self):
        """旧 chunks（没有 [IMG: ...] 标记）应保持原样不破坏。"""
        from custom_app.services.tools.list_chunks import ListChunksTool
        rows = [{
            "id": "0", "title": "T", "doc": "D",
            "contents": "纯文本，没有图片标记",
            "images": [],
        }]
        tool = ListChunksTool(rows=rows)
        result = tool.run(doc_id="D")
        assert result[0]["contents"] == "纯文本，没有图片标记"
